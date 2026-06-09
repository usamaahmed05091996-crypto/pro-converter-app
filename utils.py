import os
import pikepdf
import pdfplumber
import pandas as pd
from PIL import Image
from pdf2docx import Converter
from pypdf import PdfWriter, PdfReader
from moviepy import VideoFileClip
from typing import List

# --- CONVERSION FUNCTIONS ---

from docx import Document
import pdfplumber

def convert_pdf_to_word(pdf_path: str, docx_path: str) -> None:
    """RAM-friendly conversion: PDF to Word."""
    doc = Document()
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                doc.add_paragraph(text)
            
            # Agar table hai toh table extract karo
            table = page.extract_table()
            if table:
                t = doc.add_table(rows=len(table), cols=len(table[0]))
                for i, row in enumerate(table):
                    for j, cell in enumerate(row):
                        t.cell(i, j).text = str(cell) if cell else ""
            doc.add_page_break()
    doc.save(docx_path)

def convert_pdf_to_excel(pdf_path: str, excel_path: str) -> None:
    """Memory-efficient Excel extraction using page-by-page processing."""
    try:
        all_data = []
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                # Sirf text strategy use karein, complex lines se avoid karein
                table = page.extract_table(table_settings={
                    "vertical_strategy": "text",
                    "horizontal_strategy": "text",
                    "snap_tolerance": 3
                })
                
                if table:
                    # Header handle karein
                    df = pd.DataFrame(table[1:], columns=table[0])
                    # Duplicate columns aur empty rows hatayein
                    df = df.loc[:, ~df.columns.duplicated()]
                    df = df.dropna(how='all')
                    all_data.append(df)
        
        if not all_data:
            raise Exception("Table detect nahi ho paya, PDF ka format check karein.")
            
        final_df = pd.concat(all_data, ignore_index=True)
        final_df.to_excel(excel_path, index=False)
        
    except Exception as e:
        raise Exception(f"PDF to Excel conversion failed: {str(e)}")

def merge_pdfs(pdf_list: List[str], output_path: str) -> None:
    """Merges multiple PDF files into one."""
    try:
        writer = PdfWriter()
        for pdf_path in pdf_list:
            reader = PdfReader(pdf_path)
            for page in reader.pages:
                writer.add_page(page)
        
        with open(output_path, "wb") as f:
            writer.write(f)
    except Exception as e:
        raise Exception(f"PDF Merge failed: {str(e)}")

def protect_pdf(input_path: str, output_path: str, password: str) -> None:
    """Encrypts a PDF with a user-provided password."""
    try:
        with pikepdf.open(input_path) as pdf:
            pdf.save(output_path, encryption=pikepdf.Encryption(owner=password, user=password, R=4))
    except Exception as e:
        raise Exception(f"PDF Protection failed: {str(e)}")

def convert_image_to_pdf(image_path: str, output_pdf_path: str) -> None:
    """Converts image (JPG/PNG) to PDF format."""
    try:
        img = Image.open(image_path)
        if img.mode != 'RGB':
            img = img.convert('RGB')
        img.save(output_pdf_path, "PDF", resolution=100.0)
    except Exception as e:
        raise Exception(f"Image to PDF failed: {str(e)}")

def convert_video_to_audio(video_path: str, audio_path: str) -> None:
    """Extracts audio from video files using memory-safe closing."""
    clip = None
    try:
        clip = VideoFileClip(video_path)
        if clip.audio is None:
            raise ValueError("Video mein audio track nahi hai.")
        clip.audio.write_audiofile(audio_path, logger=None, codec='mp3')
    except Exception as e:
        raise Exception(f"Video to Audio extraction failed: {str(e)}")
    finally:
        if clip:
            clip.close()